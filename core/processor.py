from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def apply_universal_header_footer(doc):
    """
    Standardizes both headers and footers across all sections by referencing the first available (page 1) header/footer.
    Also ensures page margins and size are consistent to avoid layout issues.
    """
    if not doc.sections:
        return

    first_section = doc.sections[0]
    sectPr = first_section._sectPr
    
    # Helper to find reference ID
    def get_ref_id(sect_pr, tag_name, type_str):
        for child in sect_pr.findall(qn(tag_name)):
            if child.get(qn('w:type')) == type_str:
                return child.get(qn('r:id'))
        return None

    # Check if titlePg is active to know if we should prefer 'first'
    title_pg = sectPr.find(qn('w:titlePg'))
    has_title_pg = title_pg is not None and title_pg.get(qn('w:val')) not in ['0', 'false']

    # --- RESOLVE FOOTER ID ---
    footer_id = None
    if has_title_pg:
        footer_id = get_ref_id(sectPr, 'w:footerReference', 'first')
    if not footer_id:
        footer_id = get_ref_id(sectPr, 'w:footerReference', 'default')

    # --- RESOLVE HEADER ID ---
    header_id = None
    if has_title_pg:
        header_id = get_ref_id(sectPr, 'w:headerReference', 'first')
    if not header_id:
        header_id = get_ref_id(sectPr, 'w:headerReference', 'default')

    # --- CAPTURE MARGINS & SIZE ---
    pg_mar = sectPr.find(qn('w:pgMar'))
    pg_sz = sectPr.find(qn('w:pgSz'))

    # 2. Apply to ALL sections
    import copy
    for section in doc.sections:
        sect_pr = section._sectPr
        
        # Remove 'Different First Page' so we force standard behavior
        t_pg = sect_pr.find(qn('w:titlePg'))
        if t_pg is not None:
             sect_pr.remove(t_pg)
        
        # --- APPLY MARGINS & SIZE (Clone) ---
        if pg_mar is not None:
            curr_mar = sect_pr.find(qn('w:pgMar'))
            if curr_mar is not None:
                sect_pr.remove(curr_mar)
            sect_pr.append(copy.deepcopy(pg_mar))
            
        if pg_sz is not None:
            curr_sz = sect_pr.find(qn('w:pgSz'))
            if curr_sz is not None:
                sect_pr.remove(curr_sz)
            sect_pr.append(copy.deepcopy(pg_sz))

        # Helper to set reference
        def set_reference(tag_name, ref_id):
            if not ref_id: return
            # Clear existing refs
            for child in sect_pr.findall(qn(tag_name)):
                sect_pr.remove(child)
            
            # Add new default ref at the BEGINNING
            new_ref = OxmlElement(tag_name)
            new_ref.set(qn('w:type'), 'default')
            new_ref.set(qn('r:id'), ref_id)
            sect_pr.insert(0, new_ref)

        # Insert Footer first, then Header (Header ends up first)
        set_reference('w:footerReference', footer_id)
        set_reference('w:headerReference', header_id)


def process_with_model(content_path, model_path, output_path):
    """
    Loads a Model DOCX (with perfect headers/footers), clears its body text,
    and injects the content from content_path into it.
    """
    try:
        # Load Model (Target)
        target_doc = Document(model_path)
        
        # Load Content (Source)
        source_doc = Document(content_path)
        
        # 1. Unlock Model (just in case model itself is locked)
        # We reuse the logic from process_docx if needed, or just apply it here inline
        if target_doc.settings:
            settings_element = target_doc.settings.element
            protection = settings_element.find(qn('w:documentProtection'))
            if protection is not None:
                settings_element.remove(protection)

        # 2. Ensure Model has Universal Headers/Footers (so new pages get them)
        apply_universal_header_footer(target_doc)
        
        # 3. Clear Model Body
        # We want to keep the SECTION properties of the model (margins, headers), but remove text.
        # The body contains paragraphs and tables. We remove them.
        body_element = target_doc._element.body
        # We can't just clear() because we might lose sectPr if it's at the end? 
        # Actually sectPr is usually the last child of body. 
        # If we remove all children, we remove sectPr, and thus margins/headers.
        # We must preserve the LAST sectPr.
        
        # Strategy: Iterate all children. If it is NOT sectPr, remove it.
        for child in list(body_element):
            if child.tag.endswith('sectPr'):
                continue
            body_element.remove(child)

        # 4. Inject Content
        # We use our copy helpers (copy_paragraph, copy_table) to copy from Source to Target Body.
        # Note: We append to the body.
        # CAUTION: If we just append, we might append AFTER the sectPr if valid XML order isn't maintained (sectPr must be last).
        # We should insert BEFORE the sectPr.
        
        sect_pr = target_doc._element.body.sectPr # attribute access usually works if exists?
        # Or find the last child.
        
        # Helper to append before last sectPr
        def append_to_body(element):
            # If body has sectPr at end, insert before it. Else append.
            # python-docx's body implementation handles add_paragraph by appending.
            # If we access _element directly we must be careful.
            
            # Let's use the high-level API where possible, but we cleared the body manually.
            # If we use target_doc.add_paragraph(), it might append correctly?
            # Let's try high level first to keep it simple, but we cleared low level.
            # Re-initializing Document on the same element might be tricky.
            pass

        # Better Strategy: 
        # Use helper from before: process_container.
        # We act on target_doc.
        
        # But wait, target_doc now has no paragraphs.
        # We can just iterate source_doc and add to target_doc.
        # But we need to make sure we don't mess up the sectPr order.
        # Typically `add_paragraph()` adds before the final sectPr.
        
        for child in source_doc._element.body.iterchildren():
            # Skip sectPr from source (we want Target's layout)
            if child.tag.endswith('sectPr'):
                continue
                
            if isinstance(child, CT_P):
                p_src = Paragraph(child, source_doc)
                copy_paragraph(p_src, target_doc)
            elif isinstance(child, CT_Tbl):
                t_src = Table(child, source_doc)
                copy_table(t_src, target_doc)

        target_doc.save(output_path)
        return True, "Arquivo gerado com sucesso a partir do modelo."

    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, str(e)

def process_docx(input_path, output_path):
    """
    (Legacy/Simple Mode) - Reads a DOCX file and performs a surgical removal of protection settings.
    """
    try:
        doc = Document(input_path)
        
        # 1. Remove Document Protection
        if doc.settings:
            settings_element = doc.settings.element
            
            # Remove protection
            protection = settings_element.find(qn('w:documentProtection'))
            if protection is not None:
                settings_element.remove(protection)
                
            # Remove evenAndOddHeaders (fixes missing headers on Page 2/Even pages)
            even_odd = settings_element.find(qn('w:evenAndOddHeaders'))
            if even_odd is not None:
                settings_element.remove(even_odd)

        # 2. Standardize Footer/Header (First page -> All pages)
        apply_universal_header_footer(doc)

        # 3. Save
        doc.save(output_path)
        return True, "Arquivo processado com sucesso."
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, str(e)
